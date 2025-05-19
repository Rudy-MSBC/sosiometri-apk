
import streamlit as st
import pandas as pd
import networkx as nx
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Aplikasi Sosiometri & Sosiogram", layout="centered")

st.title("📊 Aplikasi Sosiometri & Sosiogram Berbasis Excel")

st.markdown("### 📥 Unduh Template Excel")
with open("template_sosiometri.xlsx", "rb") as f:
    st.download_button("⬇️ Download Template Excel", f, file_name="template_sosiometri.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

st.markdown("### 📤 Upload File Excel Hasil Sosiometri")
uploaded_file = st.file_uploader("Upload file Excel yang sudah diisi", type=["xlsx"])

def interpretasi_skor(skor):
    if skor == 0:
        return "❗ Terisolasi – Perlu perhatian khusus"
    elif skor <= 2:
        return "⚠️ Sosial Terbatas – Perlu dorongan interaksi"
    elif skor <= 5:
        return "✅ Cukup Sosial"
    elif skor <= 9:
        return "🌟 Populer – Disukai banyak teman"
    else:
        return "🏅 Sangat Populer – Potensial jadi fasilitator"

def proses_sosiometri(df):
    nama_siswa = df['Nama Siswa'].tolist()
    popularitas = {nama: 0 for nama in nama_siswa}
    hubungan = []
    for _, row in df.iterrows():
        pemilih = row['Nama Siswa']
        for kolom in ['Pilihan 1', 'Pilihan 2', 'Pilihan 3']:
            dipilih = row[kolom]
            if pd.notna(dipilih) and dipilih in popularitas:
                popularitas[dipilih] += 1
                hubungan.append((pemilih, dipilih))
    df_hasil = pd.DataFrame([
        {"Nama": nama, "Skor Popularitas": skor, "Interpretasi": interpretasi_skor(skor)}
        for nama, skor in popularitas.items()
    ])
    df_hasil = df_hasil.sort_values(by="Skor Popularitas", ascending=False).reset_index(drop=True)
    return df_hasil, hubungan

def simpan_word(df):
    doc = Document()
    doc.add_heading("Hasil Sosiometri dan Interpretasi", 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama'
    hdr_cells[1].text = 'Skor Popularitas'
    hdr_cells[2].text = 'Interpretasi'
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Nama'])
        row_cells[1].text = str(row['Skor Popularitas'])
        row_cells[2].text = str(row['Interpretasi'])
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def simpan_gambar_sosiogram(hubungan, df_hasil):
    G = nx.DiGraph()
    G.add_edges_from(hubungan)
    node_colors = []
    labels = {}
    for i, row in df_hasil.iterrows():
        nama = row['Nama']
        skor = row['Skor Popularitas']
        label = f"{i+1}. {nama}"
        labels[nama] = label
        if skor == 0:
            node_colors.append("red")
        elif skor <= 2:
            node_colors.append("orange")
        elif skor <= 5:
            node_colors.append("yellow")
        else:
            node_colors.append("green")
    pos = nx.spring_layout(G, seed=42)
    fig, ax = plt.subplots(figsize=(10, 8))
    nx.draw(G, pos, with_labels=False, arrows=True, node_color=node_colors, node_size=2500, edge_color='gray')
    nx.draw_networkx_labels(G, pos, labels=labels, font_size=10)
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    return buf, fig

if uploaded_file:
    df_excel = pd.read_excel(uploaded_file)
    try:
        df_hasil, hubungan = proses_sosiometri(df_excel)
        st.success("✅ Data berhasil diproses!")
        st.markdown("### 📊 Hasil Skor & Interpretasi")
        st.dataframe(df_hasil)

        st.markdown("### 🧠 Unduh Hasil Interpretasi (Word)")
        word_file = simpan_word(df_hasil)
        st.download_button("⬇️ Download Hasil Word", data=word_file, file_name="hasil_sosiometri.docx")

        st.markdown("### 🖼️ Visualisasi Sosiogram")
        png_file, fig = simpan_gambar_sosiogram(hubungan, df_hasil)
        st.pyplot(fig)
        st.download_button("⬇️ Download Gambar PNG", data=png_file, file_name="sosiogram.png", mime="image/png")
    except Exception as e:
        st.error(f"Terjadi kesalahan dalam memproses file: {e}")
